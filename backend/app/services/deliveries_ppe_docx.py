import base64
import binascii
import io
import os
from collections.abc import Iterable
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from app.models import Employee, EquipmentDelivery


LOGO_FILENAME = "logo.png"


def _maybe_get_logo_path() -> Path | None:
    candidates: list[Path] = []
    upload_dir = os.getenv("UPLOAD_DIR")
    if upload_dir:
        candidates.append(Path(upload_dir) / LOGO_FILENAME)
    candidates.append(Path(__file__).resolve().parents[2] / "upload" / LOGO_FILENAME)
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _decode_signature(signature_b64: str | None) -> io.BytesIO | None:
    if not signature_b64:
        return None
    value = signature_b64.strip()
    if not value:
        return None
    if "," in value:
        value = value.split(",", 1)[1]
    try:
        binary = base64.b64decode(value, validate=True)
    except (binascii.Error, ValueError):
        return None
    stream = io.BytesIO(binary)
    stream.seek(0)
    return stream


def export_employee_deliveries_docx(*, employee: Employee, deliveries: Iterable[EquipmentDelivery]) -> bytes:
    deliveries = list(deliveries)
    document = Document()

    logo_path = _maybe_get_logo_path()
    if logo_path:
        document.add_picture(str(logo_path), width=Inches(1.8))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading = document.add_paragraph("Scheda consegna DPI e vestiario")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if heading.runs:
        heading.runs[0].bold = True

    year_paragraph = document.add_paragraph(f"Anno {datetime.now().year}")
    year_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")
    document.add_paragraph(f"Dipendente: {employee.full_name}")
    role_value = employee.organization_role or employee.tms_role_description or "-"
    document.add_paragraph(f"Ruolo: {role_value}")
    document.add_paragraph(f"Reparto: {employee.organization_department or employee.organization_function or '-'}")

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Articolo", "Taglia", "Quantita", "Data consegna"]
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label

    for delivery in deliveries:
        row = table.add_row().cells
        row[0].text = delivery.item_name
        row[1].text = delivery.item_size or ""
        row[2].text = str(delivery.quantity)
        row[3].text = delivery.delivered_at.strftime("%d/%m/%Y")

    if len(deliveries) == 0:
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = ""
        row[3].text = ""

    document.add_paragraph("")
    document.add_paragraph("Firma del ricevente")

    signature_stream = None
    for delivery in deliveries:
        signature_stream = _decode_signature(delivery.signature_b64)
        if signature_stream:
            break
    if signature_stream:
        document.add_picture(signature_stream, width=Inches(2.8))
    else:
        document.add_paragraph("")

    footer = document.add_paragraph(
        "Il dipendente dichiara di aver ricevuto il materiale sopra indicato e di essere stato informato sul corretto utilizzo."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
